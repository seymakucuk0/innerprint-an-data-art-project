"""Generate `innerprint_manifesto.docx` — the project's manifesto,
data-annotation table, mechanics notes, and repo link in one document.
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "innerprint_manifesto.docx"

REPO_URL = "https://github.com/seymakucuk0/innerprint-an-data-art-project"


def _shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Iowan Old Style"
        run.font.color.rgb = RGBColor(0x14, 0x10, 0x1A)
        run.font.size = Pt(20 if level == 0 else 16 if level == 1 else 13)


def add_para(doc: Document, text: str, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Iowan Old Style"
    run.font.size = Pt(size)
    if italic:
        run.italic = True
    p.paragraph_format.space_after = Pt(6)


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Iowan Old Style"
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(0x3A, 0x22, 0x40)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Light Grid Accent 1"

    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        _shade(cell, "1A1020")
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h.upper())
        run.font.name = "Iowan Old Style"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xF5, 0xE9, 0xD8)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    # body rows
    for r, row in enumerate(rows, start=1):
        for c, txt in enumerate(row):
            cell = t.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.name = "Iowan Old Style"
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "3A4FB8")
    rPr.append(color)
    under = OxmlElement("w:u")
    under.set(qn("w:val"), "single")
    rPr.append(under)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Iowan Old Style")
    rPr.append(rFonts)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build() -> None:
    doc = Document()

    # margins
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── Title ─────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = title_p.add_run("innerprint")
    title.font.name = "Iowan Old Style"
    title.font.size = Pt(38)
    title.font.color.rgb = RGBColor(0x14, 0x10, 0x1A)
    title.bold = False

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = sub_p.add_run("— the spiral within —")
    sub.font.name = "Iowan Old Style"
    sub.font.size = Pt(15)
    sub.italic = True
    sub.font.color.rgb = RGBColor(0x3A, 0x22, 0x40)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = meta_p.add_run("Şeymanur Küçük  ·  2026  ·  data art installation")
    meta.font.name = "Iowan Old Style"
    meta.font.size = Pt(10)
    meta.font.color.rgb = RGBColor(0x70, 0x60, 0x70)

    doc.add_paragraph()  # spacer

    # ── Manifesto ─────────────────────────────────────────────────────
    add_heading(doc, "Manifesto", 1)

    add_quote(doc,
        "Bu bir gösterge paneli değil. İçinde yürünebilecek bir mekân."
    )

    add_para(doc,
        "innerprint, sekiz aylık bir hayatın izlerini — dinlenen müziği, "
        "atılan adımları, İstanbul'un gökyüzünü, telefona harcanan saatleri "
        "— gerçek bir parmak izinin spiral desenine dönüştürür. Tek bir "
        "giriş, tek bir çıkış. Merkez bugündür; dış kenar 2025 Ağustos'unda "
        "başlar."
    )

    add_para(doc,
        "Her parmak izi tek ve tekrar edilemez. Hayatın her sekiz ayı da "
        "öyle. Spiralin kıvrımları sıradan bir Arşimet eğrisi değildir; "
        "organik ufak salınımlarla biraz parmak izi, biraz nehir yatağı, "
        "biraz bir günün anısı olur. Yürüdükçe duvarın altında o günün "
        "müziği ısınır, üstünde o günün gökyüzü değişir, koridor o günün "
        "ekran süresine göre genişler ya da sıkıştırır. Ayağının altındaki "
        "ışık o gün ne kadar adım atıldığını söyler."
    )

    add_para(doc,
        "Sakin günler nefes alır. Sıkışmış günler donar, kararır, ara sıra "
        "siyah parıltılarla ürperir. Eser duvarlardan veri akıtmaz — "
        "veriyi mekânlaştırır."
    )

    add_quote(doc,
        "Veriyi görselleştirmiyoruz; üzerinde yürünebilir hale getiriyoruz."
    )

    # ── Conceptual frame ──────────────────────────────────────────────
    add_heading(doc, "Kavramsal Çerçeve", 1)

    add_para(doc,
        "Form, biyometrik bir parmak izinden ödünç alındı: tek girişli, "
        "merkezde sonlanan, organik salınımlı bir spiral. Her insan "
        "parmak izi kendine özgüdür; her insanın da bir dönemine ait "
        "verisi öyle. Form bu eşsizliği taşır."
    )

    add_para(doc,
        "Aesthetic referans olarak Tayvanlı dijital sanatçı Damon Xart'ın "
        "(Hsieh Chen Lin) \"gradient corridor\" çalışmaları benimsendi: "
        "sert sınırlar değil, nefes alan dikey gradyanlar; geometri "
        "üzerine dökülmüş donmuş ışık. Bu his hem heykel görünümünde "
        "(orbit modu) hem koridor içinde (walk modu) korunur."
    )

    add_para(doc,
        "Spiralin merkezi bugündür: \"bugün\" kelimesi merkez koordinatın "
        "kendisidir. Dışa doğru yürüdükçe geçmişe gidersin; içe doğru, "
        "yaşadığın güne yaklaşırsın. Bu yön kasıtlıdır: dışarısı tarih, "
        "merkez şimdidir."
    )

    # ── Data annotation table ─────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Veri Eşleme Tablosu", 1)

    add_para(doc,
        "Aşağıdaki tablo her veri akışının fiziksel/görsel karşılığını "
        "ve karar gerekçesini gösterir. 266 günün her biri için, bu "
        "kuralların kesişimi o günün duvarını üretir."
    )

    rows = [
        [
            "Spotify – günün en çok dinlenen şarkısının albümü",
            "spotify_history.csv → en yüksek play sayılı (artist, track) → albüm → signature renk",
            "Duvarın alt anchor rengi (gradyanın dibi)",
            "Müzik kişisel kimliktir; gün gün değişir. Albümün signature rengi 5 paralel AI ajan tarafından kapak görseline bakılarak seçildi — sayıca baskın değil, hatırlanan renk (Daft Punk altın, OK Computer buz-mavi, Strokes magenta).",
        ],
        [
            "Hava (sıcaklık, güneş, yağış, kar)",
            "Open-Meteo Istanbul archive — 266 gün boyunca",
            "Duvarın üst anchor rengi + tavan",
            "Sıcak güneşli gün altın amber, soğuk güneşli gün gök mavisi, kapalı gün gri-mavi, yağmurlu gri, ağır yağmurlu siyaha yakın storm, karlı buz mavisi. RGB lerp ile yumuşak geçişler.",
        ],
        [
            "Sunshine hours",
            "Open-Meteo",
            "Üst anchor parlaklığı",
            "Çok güneşli → tavan dolu ışık, az güneşli → mat. Smoothstep gating ile geçiş ne kadar açık olduğunu duvarın üstünde okutur.",
        ],
        [
            "Ekran süresi (toplam)",
            "iOS Screen Time (2026-04-13 → 2026-05-18 gerçek + öncesi 8-aylık narratif mock)",
            "Koridor genişliği + duvar yüksekliği + karanlıklaşma + glitch + siyah parlamalar",
            "Yoğun gün dar (0.92 m) + yüksek (4.9 m) + −%44 karanlık + ±%32 jitter; sakin gün ferah (3.20 m), kısa (2.8 m), nefes alan. Ekran süresi tek başına dört kanalı tetikler — bedensel sıkışma + atmosferik kararma.",
        ],
        [
            "Ekran süresi alt kategorileri (productivity / social / entertainment / other)",
            "iOS Screen Time",
            "HUD'da sol üstte mini bar chart",
            "Görsel olarak duvara binmez; bilgi katmanı olarak walking esnasında okunur. Her günün dakikalık dağılımı renkli bar olarak görünür (mavi: prod, pembe: social, amber: ent, gri: other).",
        ],
        [
            "Adım sayısı",
            "Apple Health export",
            "Zemindeki albüm rengi parlaklığı + koridor genişliğine ±%10 ek modülasyon",
            "Çok yürünen gün → yerde o günün albüm rengi ışık havuzu olarak parlar + koridor +%10 ferahlar (beden dünyada hareketliydi). Az yürünen gün → zemin neredeyse siyah + koridor −%10 sıkışır.",
        ],
        [
            "Late-night streams (gece dinlemeleri)",
            "spotify_history.csv saat bazlı",
            "(Şu an aktif değil — sonraki iterasyonda tavan karanlık derinliği)",
            "Gece dinlemeleri tavanı koyulaştıracak: \"gece çalışılmış gün\" tavandan da hissedilir. Sonraki sprintin parçası.",
        ],
    ]

    add_table(
        doc,
        headers=["Veri", "Kaynak", "Görsel Karşılığı", "Karar Gerekçesi"],
        rows=rows,
    )

    # ── Mock data note ────────────────────────────────────────────────
    add_heading(doc, "Mock Veri Politikası", 2)
    add_para(doc,
        "iOS Screen Time, takip cihazına 2026-04-13 tarihinde ilk kez "
        "eklendiğinden, öncesi 7.5 ay için gerçek veri mevcut değildi. "
        "Boş bırakmak yerine deliberate bir narratif mock üretildi: "
        "Ağustos başlangıçtan Aralık başına kadar yüksek plato (final "
        "yıl projesi yoğunluğu), Aralık-Şubat arası vadi (yarıyıl tatili, "
        "Ocak dipi), Mart-Nisan rampası (yeniden ısınma). Mock günler "
        "sanat eserine eşit ağırlıkta katılır — provenance ayrımı "
        "`data/mock_metadata.json`'da tutulur, izleyiciye gösterilmez."
    )

    # ── Mechanics ─────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Yapı ve Mekanik", 1)

    add_para(doc,
        "Spiral: 5 tur, dış yarıçap 18 m, iç yarıçap 2.5 m. Çift duvarlı "
        "koridor (2 m taban genişliği) + zemin + tavan. Walking modunda "
        "1.62 m göz yüksekliği. Heykel modunda kamera yörüngede yavaşça "
        "döner."
    )

    add_para(doc,
        "Her gün spiralin bir segmentini oluşturur; ardışık günler kesin "
        "kenarlarla değil yumuşak renk lerp'i ile birbirine bağlanır. "
        "Hiçbir veri sıçramasında sert sınır yoktur — komşu günler renk "
        "renge erir."
    )

    add_para(doc,
        "Teknik altyapı: React 18 + react-three-fiber + Three.js + drei + "
        "postprocessing (bloom). Custom GLSL shader'lar her duvarı dikey "
        "gradyan + sakin günlerde nefes/akış + yoğun günlerde karanlık + "
        "glitch katmanı olarak çizer. Spotify Web Playback SDK (PKCE auth) "
        "yürüyüş esnasında o günün şarkısını çalar — varsayılan olarak "
        "audio-analysis ile nakarat bölümünden başlar."
    )

    add_heading(doc, "Kontroller", 2)

    ctl_rows = [
        ["V", "orbit ↔ walk modunu değiştir"],
        ["↑ / W", "spiralin içine doğru zamanda ileri yürü"],
        ["↓ / S", "geriye doğru yürü"],
        ["← → / A D", "yana göz at"],
        ["J", "tarih seçici aç, doğrudan o güne atla"],
        ["P", "bulunduğun günün şarkısını nakarat'tan çal"],
        ["Shift + P", "şarkıyı en baştan çal"],
    ]
    add_table(doc, headers=["Tuş", "İşlev"], rows=ctl_rows)

    # ── Repo link ─────────────────────────────────────────────────────
    add_heading(doc, "Kaynak Kod", 1)
    add_para(doc,
        "Tüm pipeline kodu, mock üretim scriptleri, palette ajan "
        "çıktıları ve görsel katman public bir GitHub repository'sinde "
        "yayınlanır. Veri katmanı `data/processed/` altında CSV ve JSON "
        "olarak; sanat katmanı `art/` altında Vite tabanlı bir React + "
        "Three.js uygulaması olarak yer alır."
    )

    link_p = doc.add_paragraph()
    add_hyperlink(link_p, REPO_URL, REPO_URL)

    add_para(doc, "")

    add_para(doc,
        "Repository içeriği aşağıdaki bölümlere ayrılır:",
        size=11,
    )
    structure = [
        "data/processed/ — günlük birleşik tablo (innerprint_daily.csv / .json), kaynak CSV'ler, signature palet",
        "data/agent_batches/ — 5 paralel AI ajanın albüm signature renklerini seçtiği batch input ve çıktıları",
        "scripts/ — backfill_weather.py, generate_screentime.py, top_song_per_day.py, merge_innerprint.py, build_manifesto_docx.py",
        "art/ — Vite + React + Three.js prototip (canlı izlenebilir uygulama)",
        ".github/workflows/deploy.yml — GitHub Pages üzerinde otomatik dağıtım",
    ]
    for s in structure:
        bp = doc.add_paragraph(s, style="List Bullet")
        for run in bp.runs:
            run.font.name = "Iowan Old Style"
            run.font.size = Pt(10)

    # ── Footer ────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("the spiral remembers what the body did,\nwhat the ears heard, what the sky said.")
    fr.font.name = "Iowan Old Style"
    fr.font.size = Pt(11)
    fr.italic = True
    fr.font.color.rgb = RGBColor(0x3A, 0x22, 0x40)

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
